from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(4)
    section.right_margin  = Cm(3)

# ── Font helpers ──────────────────────────────────────────────────────────────
BODY_FONT = 'Times New Roman'

def set_font(run, size=12, bold=False, italic=False, color=None, font=BODY_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text='', size=12, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True,
         space_before=0, space_after=6, line_spacing=24):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(line_spacing)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p

def heading_sub(number, title, size=12):
    """Sub-chapter heading, e.g. [X].1 Pengujian Sensor pH"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(f'{number}  {title}')
    set_font(run, size=size, bold=True)
    return p

def table_caption(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f'Tabel [X].{number}  {text}')
    set_font(r, size=11, bold=True)

def figure_caption(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(f'Gambar [X].{number}  {text}')
    set_font(r, size=11, italic=True)

# ── Image placeholder (gray bordered box) ────────────────────────────────────
def figure_placeholder(label):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]

    # gray fill
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'EEEEEE')
    tcPr.append(shd)

    # fixed height hint via preferred width
    tbl.columns[0].width = Cm(12)
    cell.height = Cm(7)

    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run(f'[ {label} ]')
    set_font(r, size=11, italic=True, color=(120, 120, 120))

# ── Result table ──────────────────────────────────────────────────────────────
HEADER_COLOR = '1F4E79'
ROW_ALT      = 'DEEAF1'
ROW_WHITE    = 'FFFFFF'
AVG_COLOR    = 'BDD7EE'

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def result_table(headers, rows, avg_row, col_widths=None):
    n = len(headers)
    tbl = doc.add_table(rows=1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)

    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], HEADER_COLOR)
        cell_text(hdr[i], h, bold=True, size=10, color=(255, 255, 255))

    for idx, row in enumerate(rows):
        cells = tbl.add_row().cells
        bg = ROW_ALT if idx % 2 == 1 else ROW_WHITE
        for i, val in enumerate(row):
            shade_cell(cells[i], bg)
            cell_text(cells[i], val, size=10)

    avg_cells = tbl.add_row().cells
    for i, val in enumerate(avg_row):
        shade_cell(avg_cells[i], AVG_COLOR)
        cell_text(avg_cells[i], val, bold=True, size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ══════════════════════════════════════════════════════════════════════════════
#  [X].1  Pengujian Sensor pH
# ══════════════════════════════════════════════════════════════════════════════
heading_sub('[X].1', 'Pengujian Sensor pH')

para(
    'Pengujian sensor pH dilakukan untuk mengukur tingkat akurasi sensor pH analog yang '
    'terpasang pada sistem Smart Buoy. Sensor pH yang digunakan adalah sensor analog berbasis '
    'elektroda gel yang terhubung ke pin ADC GPIO 34 pada ESP32. Output tegangan sensor '
    'dikondisikan oleh rangkaian penguat operasional sebelum dibaca oleh ADC 12-bit ESP32.')

para(
    'Alat ukur referensi yang digunakan adalah larutan buffer pH bersertifikat dengan nilai '
    'pH 4,01 dan pH 6,86. Larutan buffer bersertifikat dipilih sebagai referensi karena nilai '
    'pH-nya telah terjamin ketepatannya oleh produsen, sehingga memberikan acuan yang valid '
    'dan terstandarisasi. Sebagai alat bantu verifikasi awal digunakan portable pen-type digital '
    'pH meter yang dilengkapi probe kaca, layar LCD digital, dan fitur kalibrasi otomatis.')

para(
    'Sebelum pengujian dimulai, sensor pH dikalibrasi menggunakan metode dua titik (two-point '
    'linear calibration) dengan larutan buffer pH 4,01 dan pH 6,86. Nilai tegangan ADC pada '
    'kedua titik kalibrasi disimpan ke memori non-volatile (NVS) ESP32 menggunakan library '
    'Preferences. Model kalibrasi membentuk persamaan regresi linier yang mengkonversi tegangan '
    'ADC menjadi nilai pH. Susunan alat pada saat pengujian ditunjukkan pada Gambar [X].1.')

figure_placeholder('Foto setup pengujian sensor pH — probe sensor + larutan buffer pH 4,01 dan 6,86')
figure_caption(1, 'Susunan Alat Pengujian Sensor pH')

para(
    'Pengujian dilakukan pada dua titik referensi yang merupakan titik kalibrasi, yaitu larutan '
    'buffer pH 4,01 dan pH 6,86. Setiap titik diukur sebanyak 5 kali, di mana setiap pengukuran '
    'merupakan rata-rata dari 15 kali pembacaan tegangan (masing-masing telah melalui trimmed '
    'mean dari 50 sampel ADC). Hasil pengujian pada larutan buffer pH 4,01 ditampilkan pada '
    'Tabel [X].1.')

table_caption(1, 'Hasil Pengujian Akurasi Sensor pH pada Larutan Buffer pH 4,01')
H_PH = ['No.', 'Referensi (pH)', 'Sensor (pH)', 'Error Absolut', 'Error Relatif (%)']
result_table(
    H_PH,
    [['1','4,01','3,77','0,2400','5,99%'],
     ['2','4,01','4,13','0,1200','2,99%'],
     ['3','4,01','3,92','0,0900','2,24%'],
     ['4','4,01','3,96','0,0500','1,25%'],
     ['5','4,01','3,84','0,1700','4,24%']],
    ['Rata-rata','4,01','3,92','0,0907','2,26%']
)

para(
    'Berdasarkan Tabel [X].1, sensor pH menghasilkan nilai rata-rata 3,92 pH pada larutan buffer '
    'pH 4,01 dengan error absolut rata-rata sebesar 0,0907 dan error relatif rata-rata 2,26%. '
    'Rentang pembacaan antara 3,77 hingga 4,13 pH menunjukkan konsistensi yang cukup baik. '
    'Hasil pengujian pada larutan buffer pH 6,86 ditampilkan pada Tabel [X].2.')

table_caption(2, 'Hasil Pengujian Akurasi Sensor pH pada Larutan Buffer pH 6,86')
result_table(
    H_PH,
    [['1','6,86','6,77','0,0900','1,31%'],
     ['2','6,86','6,66','0,2000','2,92%'],
     ['3','6,86','6,72','0,1400','2,04%'],
     ['4','6,86','6,75','0,1100','1,60%'],
     ['5','6,86','6,70','0,1600','2,33%']],
    ['Rata-rata','6,86','6,72','0,1396','2,04%']
)

para(
    'Pada pengujian dengan larutan buffer pH 6,86 (Tabel [X].2), sensor menghasilkan nilai '
    'rata-rata 6,72 pH dengan error absolut rata-rata 0,1396 dan error relatif rata-rata 2,04%. '
    'Rentang pembacaan antara 6,66 hingga 6,77 pH menunjukkan stabilitas yang baik. '
    'Error relatif di bawah 3% pada kedua titik kalibrasi mengkonfirmasi bahwa model kalibrasi '
    'dua titik yang diimplementasikan telah bekerja dengan akurat dan sensor pH layak digunakan '
    'dalam sistem monitoring kualitas air tambak udang.')

para(
    'Perlu dicatat bahwa pengujian hanya dilakukan pada rentang kalibrasi (pH 4,01—6,86) '
    'karena keterbatasan larutan buffer yang tersedia. Akurasi di luar rentang tersebut '
    '(ekstrapolasi) tidak diuji dan berpotensi menghasilkan error yang lebih besar, '
    'sesuai dengan sifat umum model kalibrasi linier.',
    space_after=10)


# ══════════════════════════════════════════════════════════════════════════════
#  [X].2  Pengujian Sensor Kekeruhan (Turbidity)
# ══════════════════════════════════════════════════════════════════════════════
heading_sub('[X].2', 'Pengujian Sensor Kekeruhan (Turbidity)')

para(
    'Sensor kekeruhan yang digunakan adalah sensor turbidity analog yang terhubung ke pin ADC '
    'GPIO 35 pada ESP32. Sensor bekerja berdasarkan prinsip transmitansi cahaya inframerah; '
    'semakin keruh air maka semakin sedikit cahaya yang diterima detektor sehingga tegangan '
    'output menurun. Kalibrasi dilakukan dengan metode dua titik menggunakan air bersih sebagai '
    'referensi 0 NTU dan sampel air keruh sebagai titik atas kalibrasi.')

para(
    'Karena tidak tersedia alat ukur turbidity standar (turbidimeter), pengujian akurasi '
    'dilakukan pada titik kalibrasi bawah yaitu 0 NTU menggunakan air bersih sebagai referensi. '
    'Pengujian pada titik 0 NTU bertujuan untuk memverifikasi bahwa kalibrasi batas bawah '
    'tersimpan dengan benar dan sensor tidak menghasilkan pembacaan positif palsu (false positive) '
    'pada kondisi air jernih. Hasil pengujian ditampilkan pada Tabel [X].3.')

table_caption(3, 'Hasil Pengujian Akurasi Sensor Kekeruhan pada Air Bersih (0 NTU)')
result_table(
    ['No.', 'Referensi (NTU)', 'Sensor (NTU)', 'Error Absolut', 'Error Relatif (%)'],
    [['1','0,0','0,0','0,0000','N/A'],
     ['2','0,0','0,0','0,0000','N/A'],
     ['3','0,0','0,0','0,0000','N/A'],
     ['4','0,0','0,0','0,0000','N/A'],
     ['5','0,0','0,0','0,0000','N/A']],
    ['Rata-rata','0,0','0,0','0,0000','N/A']
)

para(
    'Seluruh lima pengukuran menghasilkan nilai tepat 0,0 NTU dengan error absolut 0,0000. '
    'Error relatif tidak dapat dihitung karena nilai referensi bernilai nol. Hasil ini '
    'menunjukkan bahwa kalibrasi titik bawah (0 NTU) berhasil dengan baik dan sensor tidak '
    'menghasilkan pembacaan kekeruhan palsu pada kondisi air jernih.')

para(
    'Sebagai observasi kualitatif tambahan, sensor juga dicelupkan ke dalam sampel air keruh '
    '(air kopi hitam). Sensor menunjukkan respons yang signifikan dengan pembacaan NTU yang '
    'jauh lebih tinggi dibandingkan air bersih, mengkonfirmasi kemampuan sensor dalam membedakan '
    'kondisi air jernih dan air keruh. Nilai NTU pasti untuk sampel tersebut tidak dapat '
    'dikuantifikasi karena tidak tersedianya turbidimeter standar sebagai pembanding.',
    space_after=10)


# ══════════════════════════════════════════════════════════════════════════════
#  [X].3  Pengujian Sensor Suhu
# ══════════════════════════════════════════════════════════════════════════════
heading_sub('[X].3', 'Pengujian Sensor Suhu')

para(
    'Sensor suhu yang digunakan adalah DS18B20, sensor suhu digital berbasis protokol '
    'komunikasi OneWire yang terhubung ke pin GPIO 4 pada ESP32. Sensor DS18B20 merupakan '
    'sensor suhu digital yang telah terkalibrasi dari pabrik dengan akurasi tipikal ±0,5°C '
    'pada rentang suhu -10°C hingga +85°C, sehingga tidak memerlukan kalibrasi manual '
    'oleh pengguna.')

para(
    'Alat ukur referensi yang digunakan adalah TP101 Portable Digital Probe Thermometer '
    'dengan rentang pengukuran -50°C hingga +300°C, dilengkapi probe stainless steel, '
    'layar LCD digital, serta fitur HOLD dan MAX/MIN. Pengujian dilakukan dengan mencelupkan '
    'sensor DS18B20 dan termometer referensi secara bersamaan ke dalam sampel air, kemudian '
    'menunggu kedua sensor mencapai kondisi stabil sebelum pengukuran dimulai. Susunan alat '
    'pada saat pengujian ditunjukkan pada Gambar [X].2.')

figure_placeholder('Foto setup pengujian sensor suhu — probe DS18B20 + termometer TP101 dalam air')
figure_caption(2, 'Susunan Alat Pengujian Sensor Suhu')

para('Hasil pengujian sensor suhu ditampilkan pada Tabel [X].4.')

table_caption(4, 'Hasil Pengujian Akurasi Sensor Suhu DS18B20')
result_table(
    ['No.', 'Referensi (°C)', 'Sensor (°C)', 'Error Absolut (°C)', 'Error Relatif (%)'],
    [['1','28,1','27,88','0,23','0,80%'],
     ['2','28,1','27,75','0,35','1,25%'],
     ['3','28,1','27,75','0,35','1,25%'],
     ['4','28,1','27,69','0,41','1,47%'],
     ['5','28,1','27,69','0,41','1,47%']],
    ['Rata-rata','28,1','27,75','0,35','1,25%']
)

para(
    'Berdasarkan Tabel [X].4, sensor DS18B20 menghasilkan nilai rata-rata 27,75°C dengan '
    'error absolut rata-rata 0,35°C dan error relatif rata-rata 1,25%. Rentang pembacaan '
    'antara 27,69°C hingga 27,88°C menunjukkan variasi yang sangat kecil (0,19°C). '
    'Error absolut rata-rata sebesar 0,35°C berada di bawah spesifikasi akurasi pabrik '
    'DS18B20 yaitu ±0,5°C, sehingga sensor suhu dinyatakan beroperasi sesuai spesifikasi '
    'teknis dan layak digunakan dalam sistem monitoring kualitas air tambak udang.',
    space_after=10)


# ══════════════════════════════════════════════════════════════════════════════
#  [X].4  Rekapitulasi Hasil Pengujian Akurasi
# ══════════════════════════════════════════════════════════════════════════════
heading_sub('[X].4', 'Rekapitulasi Hasil Pengujian Akurasi')

para(
    'Rekapitulasi seluruh hasil pengujian akurasi sensor ditampilkan pada Tabel [X].5. '
    'Seluruh sensor yang diuji menunjukkan tingkat akurasi yang memadai untuk keperluan '
    'monitoring kualitas air tambak udang.')

table_caption(5, 'Rekapitulasi Hasil Pengujian Akurasi Seluruh Sensor')
n = 6
tbl = doc.add_table(rows=1, cols=n)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

RECAP_HDR = ['Sensor', 'Titik Uji', 'Nilai Sensor\nRata-rata', 'Error\nAbsolut', 'Error\nRelatif (%)', 'Keterangan']
hdr = tbl.rows[0].cells
for i, h in enumerate(RECAP_HDR):
    shade_cell(hdr[i], HEADER_COLOR)
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    set_font(r, size=10, bold=True, color=(255, 255, 255))
    hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

RECAP_ROWS = [
    ['pH', 'Buffer pH 4,01', '3,92 pH', '0,0907', '2,26%', 'Memenuhi syarat'],
    ['pH', 'Buffer pH 6,86', '6,72 pH', '0,1396', '2,04%', 'Memenuhi syarat'],
    ['Kekeruhan', 'Air Bersih\n(0 NTU)', '0,0 NTU', '0,0000', 'N/A', 'Memenuhi syarat'],
    ['Suhu', '28,1 °C', '27,75 °C', '0,35 °C', '1,25%', 'Memenuhi syarat'],
]
for idx, row in enumerate(RECAP_ROWS):
    cells = tbl.add_row().cells
    bg = ROW_ALT if idx % 2 == 1 else ROW_WHITE
    for i, val in enumerate(row):
        shade_cell(cells[i], bg)
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        set_font(r, size=10)
        cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_paragraph().paragraph_format.space_after = Pt(4)

para(
    'Berdasarkan Tabel [X].5, sensor pH menunjukkan error relatif di bawah 3% pada kedua '
    'titik kalibrasi (2,26% dan 2,04%), sensor kekeruhan berhasil mempertahankan pembacaan '
    '0,0 NTU pada air bersih dengan error absolut nol, dan sensor suhu DS18B20 menghasilkan '
    'error absolut 0,35°C yang berada dalam spesifikasi pabrik. Secara keseluruhan, seluruh '
    'sensor pada sistem Smart Buoy telah memenuhi kriteria akurasi yang diperlukan untuk '
    'monitoring kualitas air tambak udang.')

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/home/ediloupatty/edi/Project/project_kuliah/final_project/Pengujian_Akurasi_Sensor.docx'
doc.save(out)
print(f'Saved: {out}')
